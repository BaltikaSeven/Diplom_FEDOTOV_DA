import sys
import pandas as pd
import numpy as np
import os
import pathlib
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from PyQt5 import uic
from PyQt5.QtWidgets import *
from PyQt5.QtCore import QDate
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg
from matplotlib.figure import Figure
from scipy import stats
from sklearn.linear_model import LinearRegression
from sklearn.preprocessing import PolynomialFeatures
from sklearn.metrics import r2_score, mean_squared_error
from sklearn.pipeline import make_pipeline
import sys
import os
import pathlib

def resource_path(relative_path):
    """Возвращает корректный путь для ресурсов после упаковки в exe"""
    try:
        # PyInstaller создает временную папку и сохраняет путь в _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        ui_path = resource_path("untitled_2.ui")
        uic.loadUi(ui_path, self)
        self.setup_connections()
        self.init_variables()
        
    def setup_connections(self):
        # Страница 1
        self.radioButton_yes.toggled.connect(self.update_btn_state)
        self.btn_import_date.clicked.connect(self.import_data)
        self.btn_slice_data.clicked.connect(self.show_imported_data)
        self.btn_date_upgrade.clicked.connect(self.interpolate_data)
        self.btn_save_dates.clicked.connect(self.save_input_data)
        self.btn_date_check.clicked.connect(self.display_saved_data)
        self.btn_revork.clicked.connect(self.restart_app)
        self.btn_prepare.clicked.connect(self.prepare_data)
        self.btn_ready_date.clicked.connect(self.show_prepared_data)
        
        # Страница 2
        self.radio_first_model.toggled.connect(self.update_model_ui)
        self.radio_second_model.toggled.connect(self.update_model_ui)
        self.btn_model_create.clicked.connect(self.create_calibration_model)
        
        # Страница 3
        self.btn_save_fio.clicked.connect(self.save_fio_data)
        self.btn_save_fio_chek.clicked.connect(self.show_fio_data)
        self.btn_wave.clicked.connect(self.select_save_path)
        self.btn_otchet.clicked.connect(self.generate_report)
    
    def init_variables(self):
        self.imported_data = None
        self.vhodnie_dannie = pd.DataFrame()
        self.model_params = {}
        self.grad_table = None
        self.stvol_for_model = None
        self.ne_stvol_for_model = None
        self.current_model_data = None
        self.current_model_type = None
        self.btn_prepare.setEnabled(False)
        self.btn_ready_date.setEnabled(False)
        
        # Инициализация виджета для графиков
        self.figure = Figure(figsize=(10, 6), dpi=100)
        self.canvas = FigureCanvasQTAgg(self.figure)
        layout = QVBoxLayout()
        layout.addWidget(self.canvas)
        self.widget_window.setLayout(layout)
        
        # Инициализация для страницы 3
        self.for_otchet = pd.DataFrame(columns=[
            'fio_zakaz', 'dolznost_zakaz', 'fio_isp', 'dolznost_isp', 'beton_class'
        ])
        
        # Путь для сохранения
        self.save_path = ""
    
    # ======================= СТРАНИЦА 1 ======================= #
    def update_btn_state(self):
        is_yes = self.radioButton_yes.isChecked()
        self.btn_date_upgrade.setEnabled(is_yes and self.imported_data is not None)
        self.btn_slice_data.setEnabled(self.imported_data is not None)
        self.btn_prepare.setEnabled(self.imported_data is not None)
        self.btn_ready_date.setEnabled(
            (self.stvol_for_model is not None and is_yes) or 
            (self.ne_stvol_for_model is not None and not is_yes)
        )
        
    def import_data(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Выберите файл Excel", "", "Excel Files (*.xlsx)"
        )
        if not file_path:
            return
            
        sheet = "СТВОЛ" if self.radioButton_yes.isChecked() else "НЕ_СТВОЛ"
        try:
            self.imported_data = pd.read_excel(file_path, sheet_name=sheet)
            self.update_btn_state()
            QMessageBox.information(self, "Успех", "Данные успешно импортированы")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка импорта: {str(e)}")
    
    def show_imported_data(self):
        if self.imported_data is None:
            QMessageBox.warning(self, "Ошибка", "Нет данных для отображения")
            return
            
        dialog = QDialog(self)
        dialog.setWindowTitle("Импортированные данные")
        dialog.resize(800, 600)
        
        table = QTableWidget()
        table.setColumnCount(len(self.imported_data.columns))
        table.setRowCount(len(self.imported_data))
        table.setHorizontalHeaderLabels(self.imported_data.columns)
        
        for i in range(len(self.imported_data)):
            for j in range(len(self.imported_data.columns)):
                item = QTableWidgetItem(str(self.imported_data.iloc[i, j]))
                table.setItem(i, j, item)
                
        layout = QVBoxLayout()
        layout.addWidget(table)
        dialog.setLayout(layout)
        dialog.exec_()
    
    def interpolate_data(self):
        if self.imported_data is None:
            QMessageBox.warning(self, "Ошибка", "Сначала импортируйте данные")
            return
            
        cols = ['С_МО', 'В_МО', 'Ю_МО', 'З_МО']
        for col in cols:
            if col not in self.imported_data.columns:
                QMessageBox.critical(self, "Ошибка", f"Столбец {col} не найден в данных")
                return
                
        for idx, row in self.imported_data.iterrows():
            values = row[cols]
            if values.count() < 1: 
                continue
                
            self.imported_data.loc[idx, cols] = (
                values.interpolate(method='linear', limit_direction='both')
            )
        QMessageBox.information(self, "Успех", "Данные интерполированы")
    
    def save_input_data(self):
        data = {
            'protokol': self.Protokol_number.text(),
            'dogovor_date': self.DateEdit_1.date().toString("dd.MM.yyyy"),
            'object_name': self.object_name.toPlainText(),
            'object_adres': self.object_adres.toPlainText(),
            'oborudovanie': self.oborudovanie.toPlainText(),
            'ntd': self.ntd.toPlainText(),
            'start_date': self.DateEdit_4.date().toString("dd.MM.yyyy"),
            'end_date': self.DateEdit_5.date().toString("dd.MM.yyyy")
        }
        self.vhodnie_dannie = pd.DataFrame([data])
        QMessageBox.information(self, "Сохранено", "Входные данные успешно сохранены")
    
    def display_saved_data(self):
        if self.vhodnie_dannie.empty:
            QMessageBox.warning(self, "Внимание", "Нет сохраненных данных")
            return
            
        data = self.vhodnie_dannie.iloc[0]
        message = (
            f"Протокол № - {data['protokol']} от {data['dogovor_date']}\n"
            f"Объект: {data['object_name']}\n"
            f"Адрес: {data['object_adres']}\n"
            f"Оборудование: {data['oborudovanie']}\n"
            f"НТД: {data['ntd']}\n"
            f"Период испытаний: {data['start_date']} - {data['end_date']}"
        )
        QMessageBox.information(self, "Сохраненные данные", message)
    
    def restart_app(self):
        reply = QMessageBox.question(
            self, "Подтверждение", 
            "Перезапустить программу? Все данные будут потеряны.",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply == QMessageBox.Yes:
            self.init_variables()
            self.Protokol_number.clear()
            self.object_name.clear()
            self.object_adres.clear()
            self.oborudovanie.clear()
            self.ntd.clear()
            self.DateEdit_1.setDate(QDate.currentDate())
            self.DateEdit_4.setDate(QDate.currentDate())
            self.DateEdit_5.setDate(QDate.currentDate())
            self.radioButton_yes.setChecked(False)
            self.radioButton_no.setChecked(False)
            self.update_btn_state()
            QMessageBox.information(self, "Перезагрузка", "Программа успешно перезагружена")
    
    def prepare_data(self):
        if self.imported_data is None:
            QMessageBox.warning(self, "Ошибка", "Сначала импортируйте данные")
            return
            
        if self.radioButton_yes.isChecked():
            # Для ствола: преобразуем данные в два столбца УК и МО
            uk_values = []
            mo_values = []
            
            # Направления для сбора данных
            directions = ['С', 'В', 'Ю', 'З']
            
            for direction in directions:
                uk_col = f"{direction}_УК"
                mo_col = f"{direction}_МО"
                
                # Проверяем наличие столбцов
                if uk_col in self.imported_data.columns and mo_col in self.imported_data.columns:
                    # Добавляем пары значений, где оба значения не NaN
                    for idx, row in self.imported_data.iterrows():
                        uk_val = row[uk_col]
                        mo_val = row[mo_col]
                        if pd.notna(uk_val) and pd.notna(mo_val):
                            uk_values.append(uk_val)
                            mo_values.append(mo_val)
            
            # Создаем DataFrame из двух столбцов
            self.stvol_for_model = pd.DataFrame({
                'УК': uk_values,
                'МО': mo_values
            })
            
            msg = (f"Данные для СТВОЛ успешно подготовлены!\n"
                   f"Количество записей: {len(self.stvol_for_model)}\n"
                   f"Формат: два столбца (УК и МО)")
            QMessageBox.information(self, "Подготовка данных", msg)
            
        elif self.radioButton_no.isChecked():
            # Для не ствола: берем только 2 столбца
            cols = ['Скорость ультразвука', 'Прочность МО']
            
            # Проверяем наличие необходимых столбцов
            missing_cols = [col for col in cols if col not in self.imported_data.columns]
            if missing_cols:
                QMessageBox.critical(self, "Ошибка", 
                                    f"Отсутствуют необходимые столбцы: {', '.join(missing_cols)}")
                return
                
            # Создаем DataFrame с двумя столбцами
            self.ne_stvol_for_model = self.imported_data[cols].dropna().copy()
            # Переименовываем столбцы для единообразия
            self.ne_stvol_for_model.columns = ['УК', 'МО']
            
            msg = (f"Данные для НЕ_СТВОЛ успешно подготовлены!\n"
                   f"Количество записей: {len(self.ne_stvol_for_model)}\n"
                   f"Формат: два столбца (УК и МО)")
            QMessageBox.information(self, "Подготовка данных", msg)
        
        # Активируем кнопку просмотра подготовленных данных
        self.btn_ready_date.setEnabled(True)
        self.update_btn_state()
    
    def show_prepared_data(self):
        """Показывает всплывающее окно с подготовленными данными"""
        if self.radioButton_yes.isChecked() and self.stvol_for_model is not None:
            data = self.stvol_for_model
            title = "Подготовленные данные (СТВОЛ)"
        elif self.radioButton_no.isChecked() and self.ne_stvol_for_model is not None:
            data = self.ne_stvol_for_model
            title = "Подготовленные данные (НЕ_СТВОЛ)"
        else:
            QMessageBox.warning(self, "Ошибка", "Нет подготовленных данных для отображения")
            return
        
        dialog = QDialog(self)
        dialog.setWindowTitle(title)
        dialog.resize(800, 600)
        
        table = QTableWidget()
        table.setColumnCount(len(data.columns))
        table.setRowCount(len(data))
        table.setHorizontalHeaderLabels(data.columns)
        
        # Заполняем таблицу данными
        for i in range(len(data)):
            for j in range(len(data.columns)):
                item = QTableWidgetItem(str(data.iloc[i, j]))
                table.setItem(i, j, item)
                
        # Добавляем возможность копирования
        table.setSelectionBehavior(QAbstractItemView.SelectItems)
        table.setSelectionMode(QAbstractItemView.ContiguousSelection)
        
        layout = QVBoxLayout()
        layout.addWidget(table)
        
        # Кнопка закрытия
        btn_close = QPushButton("Закрыть")
        btn_close.clicked.connect(dialog.accept)
        layout.addWidget(btn_close)
        
        dialog.setLayout(layout)
        dialog.exec_()
    
    # ======================= СТРАНИЦА 2 ======================= #
    def update_model_ui(self):
        """Обновляет состояние UI в зависимости от выбранной модели"""
        if self.radio_first_model.isChecked():
            self.radio_second_model.setChecked(False)
        elif self.radio_second_model.isChecked():
            self.radio_first_model.setChecked(False)
    
    def create_calibration_model(self):
        """Создает градуировочную зависимость в зависимости от выбранного типа"""
        # Определяем, какие данные использовать
        if self.radioButton_yes.isChecked() and self.stvol_for_model is not None:
            data = self.stvol_for_model
        elif self.radioButton_no.isChecked() and self.ne_stvol_for_model is not None:
            data = self.ne_stvol_for_model
        else:
            QMessageBox.warning(self, "Ошибка", "Нет подготовленных данных для построения модели")
            return
        
        # Сохраняем текущие данные для использования
        self.current_model_data = data
        
        # Очищаем предыдущий график
        self.figure.clear()
        ax = self.figure.add_subplot(111)
        
        # Извлекаем данные
        x = data['УК'].values
        y = data['МО'].values
        
        # Проверяем, достаточно ли данных
        if len(x) < 3:
            QMessageBox.warning(self, "Ошибка", "Недостаточно данных для построения модели")
            return
        
        # Вычисляем статистики
        min_x, max_x = np.min(x), np.max(x)
        min_y, max_y = np.min(y), np.max(y)
        mean_x, mean_y = np.mean(x), np.mean(y)
        
        # Определяем тип модели
        if self.radio_first_model.isChecked():
            # Линейная модель (статистические методы)
            self.current_model_type = 'linear'
            result_text = self.create_linear_model(ax, x, y, min_x, max_x, min_y, max_y, mean_x, mean_y)
        elif self.radio_second_model.isChecked():
            # Квадратичная модель (методы идентификации)
            self.current_model_type = 'quadratic'
            result_text = self.create_quadratic_model(ax, x, y, min_x, max_x, min_y, max_y, mean_x, mean_y)
        else:
            QMessageBox.warning(self, "Ошибка", "Выберите тип модели")
            return
        
        # Настраиваем график
        ax.set_xlabel('Скорость ультразвука (УК)')
        ax.set_ylabel('Прочность (МО)')
        
        # Устанавливаем светлый стиль с сеткой
        ax.grid(True, linestyle='--', alpha=0.7)
        ax.set_facecolor('white')
        self.figure.patch.set_facecolor('white')
        
        # Добавляем легенду
        ax.legend()
        
        # Обновляем canvas
        self.canvas.draw()
        
        # Выводим результаты в текстовое поле
        self.kf_obj.setPlainText(result_text)
    
    def create_linear_model(self, ax, x, y, min_x, max_x, min_y, max_y, mean_x, mean_y):
        """Создает линейную градуировочную зависимость"""
        # Преобразуем данные для модели
        X = x.reshape(-1, 1)
        
        # Создаем и обучаем модель
        model = LinearRegression()
        model.fit(X, y)
        
        # Получаем предсказания
        y_pred = model.predict(X)
        
        # Рассчитываем коэффициенты
        a = model.coef_[0]
        b = model.intercept_
        
        # Рассчитываем статистики
        n = len(x)
        residuals = y - y_pred
        s = np.sqrt(np.sum(residuals**2) / (n - 2))  # Стандартное отклонение остатков
        r = np.corrcoef(x, y)[0, 1]  # Коэффициент корреляции
        s_r = s / mean_y  # Относительная ошибка
        
        # Проверка условий ГОСТ
        gost_check = "Соответствует ГОСТ 17624-2021" if (s_r < 0.15 and r > 0.7) else "Не соответствует ГОСТ 17624-2021"
        
        # Рассчитываем Rн (среднее расчетное значение)
        R_n = np.mean(y_pred)
        
        # Формируем текст формулы
        if b >= 0:
            formula = f"R = {a:.4f} * H + {b:.4f}"
        else:
            formula = f"R = {a:.4f} * H - {abs(b):.4f}"
        
        # Строим график
        x_plot = np.linspace(min_x - 0.1, max_x + 0.1, 100)
        y_plot = model.predict(x_plot.reshape(-1, 1))
        
        # Основная линия регрессии
        ax.plot(x_plot, y_plot, 'b-', linewidth=2, label='Градуировочная зависимость')
        
        # Рассеяние исходных точек
        ax.scatter(x, y, c='r', alpha=0.7, label='Экспериментальные данные')
        
        # Доверительные интервалы (95%)
        se = np.sqrt(np.sum(residuals**2) / (n - 2)) * np.sqrt(1/n + (x_plot - mean_x)**2 / np.sum((x - mean_x)**2))
        t_val = stats.t.ppf(0.975, n - 2)  # t-значение для 95% доверительного интервала
        
        ax.fill_between(x_plot, y_plot - t_val * se, y_plot + t_val * se, 
                        color='blue', alpha=0.2, label='Доверительный интервал (95%)')
        
        # Формируем текст результатов
        result_text = f"Линейная градуировочная зависимость:\n\n"
        result_text += f"Формула: {formula}\n\n"
        result_text += "Коэффициенты:\n"
        result_text += f"a (угловой коэффициент) = {a:.4f}\n"
        result_text += f"b (свободный член) = {b:.4f}\n"
        result_text += f"S (стандартное отклонение остатков) = {s:.4f}\n"
        result_text += f"r (коэффициент корреляции) = {r:.4f}\n"
        result_text += f"S/R (относительная ошибка) = {s_r:.4f}\n\n"
        result_text += f"Проверка условий ГОСТ:\n"
        result_text += f"- S/R < 0.15: {'Да' if s_r < 0.15 else 'Нет'} ({s_r:.4f} {'<' if s_r < 0.15 else '>='} 0.15)\n"
        result_text += f"- r > 0.7: {'Да' if r > 0.7 else 'Нет'} ({r:.4f} {'>' if r > 0.7 else '<='} 0.7)\n"
        result_text += f"Итог: {gost_check}\n\n"
        result_text += "Основные характеристики:\n"
        result_text += f"Ri,min = {min_y:.2f}\n"
        result_text += f"Ri,max = {max_y:.2f}\n"
        result_text += f"Hi,min = {min_x:.2f}\n"
        result_text += f"Hi,max = {max_x:.2f}\n"
        result_text += f"Среднее Rф = {mean_y:.2f}\n"
        result_text += f"Среднее H = {mean_x:.2f}\n"
        result_text += f"Среднее Rн = {R_n:.2f}"
        
        return result_text
    
    def create_quadratic_model(self, ax, x, y, min_x, max_x, min_y, max_y, mean_x, mean_y):
        """Создает квадратичную градуировочную зависимость"""
        # Преобразуем данные для модели
        X = x.reshape(-1, 1)
        
        # Создаем и обучаем модель
        model = make_pipeline(PolynomialFeatures(2), LinearRegression())
        model.fit(X, y)
        
        # Получаем предсказания
        y_pred = model.predict(X)
        
        # Извлекаем коэффициенты
        coefs = model.named_steps['linearregression'].coef_
        intercept = model.named_steps['linearregression'].intercept_
        beta0 = intercept
        beta1 = coefs[1] if len(coefs) > 1 else 0
        beta2 = coefs[2] if len(coefs) > 2 else 0
        
        # Рассчитываем статистики
        n = len(x)
        r2 = r2_score(y, y_pred)  # Коэффициент детерминации
        
        # Скорректированный R²
        adj_r2 = 1 - (1 - r2) * (n - 1) / (n - 3)  # p=2 (квадратичная модель)
        
        # Стандартная ошибка оценки (SEE)
        see = np.sqrt(np.sum((y - y_pred)**2) / (n - 3))
        
        # F-статистика
        tss = np.sum((y - np.mean(y))**2)  # Total sum of squares
        rss = np.sum((y - y_pred)**2)  # Residual sum of squares
        f_stat = ((tss - rss) / 2) / (rss / (n - 3))
        
        # Формируем текст формулы
        formula = f"R = {beta2:.4f} * H²"
        if beta1 >= 0:
            formula += f" + {beta1:.4f} * H"
        else:
            formula += f" - {abs(beta1):.4f} * H"
            
        if beta0 >= 0:
            formula += f" + {beta0:.4f}"
        else:
            formula += f" - {abs(beta0):.4f}"
        
        # Строим график
        x_plot = np.linspace(min_x - 0.1, max_x + 0.1, 100)
        y_plot = model.predict(x_plot.reshape(-1, 1))
        
        # Основная линия регрессии
        ax.plot(x_plot, y_plot, 'b-', linewidth=2, label='Градуировочная зависимость')
        
        # Рассеяние исходных точек
        ax.scatter(x, y, c='r', alpha=0.7, label='Экспериментальные данные')
        
        # Формируем текст результатов
        result_text = f"Квадратичная градуировочная зависимость:\n\n"
        result_text += f"Формула: {formula}\n\n"
        result_text += "Коэффициенты регрессии:\n"
        result_text += f"β0 (свободный член) = {beta0:.4f}\n"
        result_text += f"β1 (линейный коэффициент) = {beta1:.4f}\n"
        result_text += f"β2 (квадратичный коэффициент) = {beta2:.4f}\n\n"
        result_text += "Статистические показатели:\n"
        result_text += f"Коэффициент детерминации (R²) = {r2:.4f}\n"
        result_text += f"Скорректированный R² = {adj_r2:.4f}\n"
        result_text += f"Стандартная ошибка оценки (SEE) = {see:.4f}\n"
        result_text += f"F-статистика = {f_stat:.4f}\n\n"
        result_text += "Основные характеристики:\n"
        result_text += f"Ri,min = {min_y:.2f}\n"
        result_text += f"Ri,max = {max_y:.2f}\n"
        result_text += f"Hi,min = {min_x:.2f}\n"
        result_text += f"Hi,max = {max_x:.2f}\n"
        result_text += f"Среднее Rф = {mean_y:.2f}\n"
        result_text += f"Среднее H = {mean_x:.2f}"
        
        return result_text
    
    # ======================= СТРАНИЦА 3 ======================= #
    def save_fio_data(self):
        """Сохраняет данные представителей в датафрейм for_otchet"""
        # Собираем данные из полей ввода
        fio_zakaz = self.fio_zakaz.text().strip()
        dolznost_zakaz = self.dolznost_zakaz.text().strip()
        fio_isp = self.fio_isp.text().strip()
        dolznost_isp = self.dolznost_isp.text().strip()
        beton_class = self.beton_class.text().strip()
        
        # Проверяем заполненность обязательных полей
        if not fio_zakaz or not dolznost_zakaz or not fio_isp or not dolznost_isp:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, заполните все поля представителей")
            return
            
        # Создаем новую запись
        new_data = {
            'fio_zakaz': fio_zakaz,
            'dolznost_zakaz': dolznost_zakaz,
            'fio_isp': fio_isp,
            'dolznost_isp': dolznost_isp,
            'beton_class': beton_class
        }
        
        # Создаем новый DataFrame с одной строкой
        new_df = pd.DataFrame([new_data])
        
        # Обновляем основной датафрейм
        self.for_otchet = new_df
        
        # Оповещаем пользователя
        QMessageBox.information(self, "Сохранено", "Данные представителей успешно сохранены")
        
        # Очищаем поля ввода (опционально)
        # self.clear_fio_fields()
    
    def show_fio_data(self):
        """Показывает сохраненные данные представителей"""
        if self.for_otchet.empty:
            QMessageBox.warning(self, "Предупреждение", "Нет сохраненных данных о представителях")
            return
            
        # Берем последнюю сохраненную запись
        data = self.for_otchet.iloc[0]
        
        # Форматируем сообщение
        message = (
            f"Представитель заказчика: {data['fio_zakaz']}\n"
            f"Должность: {data['dolznost_zakaz']}\n\n"
            f"Представитель исполнителя: {data['fio_isp']}\n"
            f"Должность: {data['dolznost_isp']}\n\n"
        )
        
        # Добавляем класс бетона, если он указан
        if pd.notna(data['beton_class']) and data['beton_class'].strip():
            message += f"Класс бетона исследуемого объекта: {data['beton_class']}"
        
        # Создаем диалоговое окно с улучшенным форматированием
        dialog = QDialog(self)
        dialog.setWindowTitle("Сохраненные данные представителей")
        dialog.setMinimumWidth(400)
        
        layout = QVBoxLayout()
        
        # Создаем текстовый элемент с HTML-форматированием
        text_edit = QTextEdit()
        text_edit.setReadOnly(True)
        text_edit.setHtml(
            f"<h3>Данные представителей</h3>"
            f"<p><b>Представитель заказчика:</b> {data['fio_zakaz']}</p>"
            f"<p><b>Должность:</b> {data['dolznost_zakaz']}</p>"
            f"<p><b>Представитель исполнителя:</b> {data['fio_isp']}</p>"
            f"<p><b>Должность:</b> {data['dolznost_isp']}</p>"
        )
        
        # Добавляем класс бетона, если он указан
        if pd.notna(data['beton_class']) and data['beton_class'].strip():
            text_edit.append(f"<p><b>Класс бетона исследуемого объекта:</b> {data['beton_class']}</p>")
        
        layout.addWidget(text_edit)
        
        # Кнопка закрытия
        btn_close = QPushButton("Закрыть")
        btn_close.clicked.connect(dialog.accept)
        layout.addWidget(btn_close)
        
        dialog.setLayout(layout)
        dialog.exec_()
    
    def clear_fio_fields(self):
        """Очищает поля ввода на странице 3 (опционально)"""
        self.fio_zakaz.clear()
        self.dolznost_zakaz.clear()
        self.fio_isp.clear()
        self.dolznost_isp.clear()
        self.beton_class.clear()
        
    def select_save_path(self):
        """Выбирает путь для сохранения отчета"""
        folder_path = QFileDialog.getExistingDirectory(
            self, "Выберите папку для сохранения отчета"
        )
        if folder_path:
            self.save_path = folder_path
            QMessageBox.information(
                self, "Путь сохранения", 
                f"Отчеты будут сохраняться в папку:\n{self.save_path}"
            )
        else:
            QMessageBox.warning(self, "Ошибка", "Папка не выбрана")
    
    def generate_report(self):
        """Генерирует отчет в формате .docx"""
        # Проверяем, выбран ли путь сохранения
        if not self.save_path:
            QMessageBox.warning(self, "Ошибка", "Сначала выберите папку для сохранения")
            return
            
        # Проверяем наличие необходимых данных
        if self.vhodnie_dannie.empty:
            QMessageBox.warning(self, "Ошибка", "Нет данных протокола")
            return
            
        if self.for_otchet.empty:
            QMessageBox.warning(self, "Ошибка", "Нет данных представителей")
            return
            
        if self.current_model_data is None or self.current_model_type is None:
            QMessageBox.warning(self, "Ошибка", "Не построена градуировочная зависимость")
            return
            
        try:
            # Создаем документ
            doc = Document()
            
            # Устанавливаем шрифт Aptos Narrow для всего документа
            self.set_aptos_font(doc)
            
            # Добавляем заголовок протокола
            self.add_protocol_header(doc)
            
            # Добавляем основное содержание отчета
            self.add_report_content(doc)
            
            # Сохраняем документ
            prot_data = self.vhodnie_dannie.iloc[0]
            report_name = f"Приложение к протоколу {prot_data['protokol']}.docx"
            report_path = os.path.join(self.save_path, report_name)
            doc.save(report_path)
            
            QMessageBox.information(
                self, "Отчет создан", 
                f"Отчет успешно сохранен по пути:\n{report_path}"
            )
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка при создании отчета: {str(e)}")
    
    def set_aptos_font(self, doc):
        """Устанавливает шрифт Aptos Narrow для всего документа"""
        # Для основного текста
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Aptos Narrow'
        font.size = Pt(11)
        
        # Для заголовков
        for heading in ['Heading 1', 'Heading 2', 'Heading 3']:
            if heading in doc.styles:
                style = doc.styles[heading]
                font = style.font
                font.name = 'Aptos Narrow'
                font.size = Pt(11)
    
    def add_protocol_header(self, doc):
        """Добавляет заголовок протокола"""
        prot_data = self.vhodnie_dannie.iloc[0]
        
        # Заголовок
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Приложение к протоколу № {prot_data['protokol']} от {prot_data['dogovor_date']}")
        run.bold = True
        run.font.size = Pt(12)
        
        # Пустая строка
        doc.add_paragraph()
        
        # Информация об объекте
        p = doc.add_paragraph()
        p.add_run(f"Объект: {prot_data['object_name']}")
        
        p = doc.add_paragraph()
        p.add_run(f"Адрес: {prot_data['object_adres']}")
        
        p = doc.add_paragraph()
        p.add_run(f"Оборудование: {prot_data['oborudovanie']}")
        
        p = doc.add_paragraph()
        p.add_run(f"НТД: {prot_data['ntd']}")
        
        p = doc.add_paragraph()
        p.add_run(f"Период испытаний: {prot_data['start_date']} - {prot_data['end_date']}")
        
        # Пустая строка
        doc.add_paragraph()
    
    def add_report_content(self, doc):
        """Добавляет основное содержание отчета"""
        # Заголовок раздела
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Построение градуировочной зависимости")
        run.bold = True
        run.font.size = Pt(12)
        
        # Пустая строка
        doc.add_paragraph()
        
        # Создаем таблицу для данных
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Устанавливаем ширину столбцов (в дюймах)
        widths = (Inches(0.5), Inches(1.5), Inches(2.0), Inches(2.0), Inches(2.0))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
        
        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№'
        hdr_cells[1].text = 'Номер участка'
        hdr_cells[2].text = 'Прочность бетона по результатам испытаний (Riф, МПа)'
        hdr_cells[3].text = 'Скорость ультразвука (Нi, м/с)'
        hdr_cells[4].text = 'Прочность бетона по градуировочной зависимости (RH, МПа)'
        
        # Жирный шрифт для заголовков
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Заполняем таблицу данными
        data = self.current_model_data
        for i, row in data.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(i+1)
            row_cells[1].text = f"Участок {i+1}"  # Номер участка
            row_cells[2].text = f"{row['МО']:.2f}"
            row_cells[3].text = f"{row['УК']:.2f}"
            
            # Предсказанная прочность
            if self.current_model_type == 'linear' and hasattr(self, 'current_linear_model'):
                rh = self.current_linear_model.predict([[row['УК']]])[0]
                row_cells[4].text = f"{rh:.2f}"
            elif self.current_model_type == 'quadratic' and hasattr(self, 'current_quadratic_model'):
                rh = self.current_quadratic_model.predict([[row['УК']]])[0]
                row_cells[4].text = f"{rh:.2f}"
            else:
                row_cells[4].text = "—"
        
        # Добавляем формулу зависимости
        doc.add_paragraph()
        p = doc.add_paragraph()
        if self.current_model_type == 'linear' and hasattr(self, 'linear_formula'):
            run = p.add_run(f"Градуировочная зависимость: {self.linear_formula}")
        elif self.current_model_type == 'quadratic' and hasattr(self, 'quadratic_formula'):
            run = p.add_run(f"Градуировочная зависимость: {self.quadratic_formula}")
        else:
            run = p.add_run("Градуировочная зависимость: не определена")
        run.bold = True
        
        # Сохраняем график во временный файл и добавляем в документ
        temp_img_path = os.path.join(self.save_path, "temp_plot.png")
        self.figure.savefig(temp_img_path, dpi=300, bbox_inches='tight')
        doc.add_picture(temp_img_path, width=Inches(4.0))
        
        # Центрируем изображение
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавляем статистические данные
        self.add_statistics_section(doc)
        
        # Добавляем данные представителей
        self.add_representatives_section(doc)
        
        # Удаляем временный файл
        if os.path.exists(temp_img_path):
            os.remove(temp_img_path)
    
    def add_statistics_section(self, doc):
        """Добавляет раздел со статистическими данными"""
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Статистические характеристики:")
        run.bold = True
        
        # Создаем таблицу для статистик
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Устанавливаем ширину столбцов
        widths = (Inches(1.5), Inches(1.5), Inches(2.0), Inches(2.0))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
        
        # Заголовки
        header_cells = table.rows[0].cells
        header_cells[0].text = "Параметр"
        header_cells[1].text = "Значение"
        header_cells[2].text = "Характеристика"
        header_cells[3].text = "Значение"
        
        # Данные
        data = self.current_model_data
        data_cells = table.rows[1].cells
        data_cells[0].text = "Среднее Rф:"
        data_cells[1].text = f"{data['МО'].mean():.2f}"
        data_cells[2].text = "Ri,min:"
        data_cells[3].text = f"{data['МО'].min():.2f}"
        
        # Добавляем вторую строку с данными
        row_cells = table.add_row().cells
        row_cells[0].text = "Среднее H:"
        row_cells[1].text = f"{data['УК'].mean():.2f}"
        row_cells[2].text = "Ri,max:"
        row_cells[3].text = f"{data['МО'].max():.2f}"
        
        # Третья строка
        row_cells = table.add_row().cells
        row_cells[0].text = "Среднее RH:"
        
        # Вычисляем среднее RH
        if self.current_model_type == 'linear' and hasattr(self, 'current_linear_model'):
            rh = self.current_linear_model.predict(data[['УК']].values.reshape(-1, 1))
            row_cells[1].text = f"{rh.mean():.2f}"
        elif self.current_model_type == 'quadratic' and hasattr(self, 'current_quadratic_model'):
            rh = self.current_quadratic_model.predict(data[['УК']].values.reshape(-1, 1))
            row_cells[1].text = f"{rh.mean():.2f}"
        else:
            row_cells[1].text = "—"
            
        row_cells[2].text = "Hi,min:"
        row_cells[3].text = f"{data['УК'].min():.2f}"
        
        # Четвертая строка
        row_cells = table.add_row().cells
        row_cells[2].text = "Hi,max:"
        row_cells[3].text = f"{data['УК'].max():.2f}"
    
    def add_representatives_section(self, doc):
        """Добавляет раздел с данными представителей"""
        if self.for_otchet.empty:
            return
            
        rep_data = self.for_otchet.iloc[0]
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Представители:")
        run.bold = True
        
        # Таблица представителей
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Устанавливаем ширину столбцов
        widths = (Inches(2.0), Inches(4.0))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
        
        # Заполняем таблицу
        cells = table.rows[0].cells
        cells[0].text = "Представитель заказчика:"
        cells[1].text = rep_data['fio_zakaz']
        
        cells = table.rows[1].cells
        cells[0].text = "Должность:"
        cells[1].text = rep_data['dolznost_zakaz']
        
        cells = table.rows[2].cells
        cells[0].text = "Представитель исполнителя:"
        cells[1].text = rep_data['fio_isp']
        
        # Добавляем строку для должности исполнителя
        row_cells = table.add_row().cells
        row_cells[0].text = "Должность:"
        row_cells[1].text = rep_data['dolznost_isp']
        
        # Добавляем класс бетона, если указан
        if rep_data['beton_class']:
            row_cells = table.add_row().cells
            row_cells[0].text = "Класс бетона:"
            row_cells[1].text = rep_data['beton_class']

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())